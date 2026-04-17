from io import BytesIO
from pypdf import PdfReader
from docx import Document


def extract_pdf(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    return "\n\n".join(page.extract_text() or "" for page in reader.pages)


def extract_docx(data: bytes) -> str:
    doc = Document(BytesIO(data))
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract(filename: str, data: bytes) -> str:
    name = filename.lower()
    if name.endswith(".pdf"):
        return extract_pdf(data)
    if name.endswith(".docx"):
        return extract_docx(data)
    raise ValueError(f"Unsupported file type: {filename}")
